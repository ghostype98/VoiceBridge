# -*- coding: utf-8 -*-
"""
根据部署文档.md、用户手册.md 的完整内容，逐段生成格式正确的 Word 文档，内容与 md 一致。
"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color="E7E6E6"):
    """表格单元格底纹"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def add_para(doc, text, bold_parts=None):
    """添加段落，可选部分加粗。bold_parts 为需加粗的字符串列表。"""
    p = doc.add_paragraph()
    if not bold_parts:
        p.add_run(text)
        return p
    remain = text
    for b in bold_parts:
        if b in remain:
            idx = remain.find(b)
            if idx > 0:
                p.add_run(remain[:idx])
            r = p.add_run(b)
            r.bold = True
            remain = remain[idx + len(b):]
    if remain:
        p.add_run(remain)
    return p

def add_code(doc, code_text):
    """添加等宽代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(code_text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_table(doc, headers, rows, first_row_bold=True):
    """添加表格：headers 为表头列表，rows 为数据行列表的列表。"""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        if first_row_bold:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
        set_cell_shading(cell)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            if ci < len(t.rows[ri + 1].cells):
                t.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()
    return t

def build_部署文档():
    doc = Document()
    # 标题
    doc.add_heading('VoiceBridge 部署文档', 0)
    add_para(doc, '本文档说明如何在服务器或本机从零部署 VoiceBridge 语音交互服务，每一步均配有简要说明；包含生产环境开机自启（systemd）配置。', ['包含生产环境开机自启（systemd）'])
    doc.add_paragraph()
    doc.add_paragraph('─' * 40)
    # 一、环境要求
    doc.add_heading('一、环境要求', 1)
    add_table(doc, ['项目', '要求'], [
        ['操作系统', 'Linux（推荐 Ubuntu 20.04+）或 macOS'],
        ['Python', '3.8 及以上'],
        ['Conda', 'Anaconda3 或 Miniconda3（用于管理 Python 环境）'],
        ['端口', '默认 8002（主服务）、8003（WebSocket），以 config/config.yaml 为准，确保未被占用'],
    ])
    add_para(doc, '说明：项目通过 Conda 使用名为 datastore 的环境运行，请先安装 Conda。', ['说明'])
    doc.add_paragraph()
    doc.add_paragraph('─' * 40)
    # 二、获取项目代码
    doc.add_heading('二、第一步：获取项目代码', 1)
    add_para(doc, '将项目克隆或拷贝到目标目录，例如：')
    doc.add_paragraph()
    add_code(doc, '# 使用 Git 克隆（仓库地址）\ngit clone https://github.com/YOUR_ORG/voicebridge.git /opt/voicebridge\ncd /opt/voicebridge')
    add_para(doc, '说明：仓库地址为 https://github.com/YOUR_ORG/voicebridge.git；可替换为实际部署目录。后续所有脚本路径均相对于项目根目录（即包含 app、bash、config 的目录）。', ['说明'])
    doc.add_paragraph()
    doc.add_paragraph('─' * 40)
    # 三、Conda 环境
    doc.add_heading('三、第二步：创建并激活 Conda 环境', 1)
    add_para(doc, '项目使用 datastore 环境。若尚未创建，可先创建再安装依赖；若已存在则直接激活。', ['datastore'])
    doc.add_paragraph()
    add_code(doc, '# 创建环境（仅首次需要）\nconda create -n datastore python=3.10 -y\n\n# 激活环境（每次新开终端都需要执行）\nconda activate datastore')
    add_para(doc, '说明：conda activate datastore 后，终端提示符前会显示 (datastore)，表示当前已在该环境中。', ['说明'])
    doc.add_paragraph()
    doc.add_paragraph('─' * 40)
    # 四、安装依赖
    doc.add_heading('四、第三步：安装 Python 依赖', 1)
    add_para(doc, '在项目根目录下、且已激活 datastore 环境时执行：')
    doc.add_paragraph()
    add_code(doc, 'cd /opt/voicebridge\npip install -r requirements.txt')
    add_para(doc, '说明：')
    doc.add_paragraph('requirements.txt 包含 FastAPI、uvicorn、阿里云 SDK、WebSocket、日志等依赖。', style='List Bullet')
    doc.add_paragraph('若提示权限错误，可加 --user 或使用虚拟环境。', style='List Bullet')
    doc.add_paragraph('安装完成后可用 python -c "import fastapi, uvicorn; print(\'ok\')" 快速验证。', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('─' * 40)
    # 五、配置文件
    doc.add_heading('五、第四步：配置文件（端口唯一修改处）', 1)
    add_para(doc, '端口与核心服务配置均只需修改一个文件：config/config.yaml。修改该文件中的端口后，重启服务即可；启动脚本与 NATAPP 会自动从该文件读取主端口。', ['config/config.yaml'])
    add_para(doc, '主配置文件为 config/config.yaml。部署时需要至少确认以下部分。', ['config/config.yaml'])
    doc.add_heading('5.1 语音流式服务（ASR）', 2)
    doc.add_paragraph('路径：config/config.yaml → voice_streaming.asr', style='List Bullet')
    doc.add_paragraph('必填：阿里云 ASR 的 appkey、access_key_id、access_key_secret', style='List Bullet')
    doc.add_paragraph('说明：用于实时语音转文字。可从阿里云控制台创建/获取；敏感信息也可通过环境变量或 config/credentials.yaml 配置（见 auto_config 段）。', style='List Bullet')
    doc.add_heading('5.2 主服务端口（唯一端口配置处）', 2)
    doc.add_paragraph('路径：config/config.yaml → services.voicebridge.port 与 server.port（二者需一致）', style='List Bullet')
    doc.add_paragraph('默认：host: "0.0.0.0"、port: 8002', style='List Bullet')
    doc.add_paragraph('说明：主服务端口同时提供 HTTP API 与前端静态页面，无需单独起前端服务。WebSocket 端口见 voice_streaming.websocket.port（默认 8003）。以后更换端口只需改本文件上述三处。', style='List Bullet')
    doc.add_heading('5.3 对话与 LLM（可选）', 2)
    doc.add_paragraph('Rasa：rasa.endpoint、rasa.port（默认 8012），若不需要对话管理可设 enabled: false。', style='List Bullet')
    doc.add_paragraph('LLM：llm.api_base、llm.model 等，按实际使用的模型服务填写。', style='List Bullet')
    doc.add_heading('5.4 日志与存储', 2)
    doc.add_paragraph('日志会写入 logs/ 下各子目录（如 logs/service/、logs/app/），无需在 config 里改即可使用。', style='List Bullet')
    doc.add_paragraph('存储目录见 storage.path、storage.audio_path，默认为 ./storage、./storage/audio。', style='List Bullet')
    add_para(doc, '说明：修改配置后需重启服务生效。', ['说明'])
    doc.add_paragraph()
    doc.add_paragraph('─' * 40)
    # 六、前端
    doc.add_heading('六、第五步：前端静态资源（可选）', 1)
    add_para(doc, '若项目根目录下存在 frontend 目录，主服务会自动在主服务端口（默认 8002）提供登录、面试等页面（/、/login、/interview 等）。', ['frontend'])
    doc.add_paragraph('若没有 frontend 目录，仅 API 可用，前端页面不可用。', style='List Bullet')
    doc.add_paragraph('若有，无需再单独启动前端开发服务器，直接访问 http://<主机>:8002 即可（端口以 config 为准）。', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('─' * 40)
    # 七、启动服务
    doc.add_heading('七、第六步：启动服务', 1)
    add_para(doc, '在项目根目录执行：')
    doc.add_paragraph()
    add_code(doc, 'bash /opt/voicebridge/bash/start_service.sh')
    add_para(doc, '脚本会依次执行：', ['脚本会依次执行'])
    for i, s in enumerate([
        '解析项目根目录与日志目录。',
        '自动查找并 source Conda（如 ~/anaconda3/etc/profile.d/conda.sh 或 miniconda 路径）。',
        '执行 conda activate datastore。',
        '设置 PYTHONPATH 为项目根目录。',
        '检查 voicebridge.pid：若已有对应进程在跑则提示“服务已在运行”并退出。',
        '使用 nohup python services/run.py 后台启动服务，并将 PID 写入 voicebridge.pid。',
        '等待约 3 秒后检查进程是否存在，并输出访问地址与日志路径。',
    ], 1):
        doc.add_paragraph(f'{i}. {s}', style='List Number')
    add_para(doc, '成功时终端会看到类似（端口以 config/config.yaml 为准，默认 8002/8003）：', ['成功时终端会看到类似'])
    doc.add_paragraph('主服务: http://0.0.0.0:8002', style='List Bullet')
    doc.add_paragraph('WebSocket: ws://0.0.0.0:8003', style='List Bullet')
    doc.add_paragraph('前端: http://localhost:8002/login', style='List Bullet')
    doc.add_paragraph('API 文档: http://localhost:8002/docs', style='List Bullet')
    doc.add_paragraph('日志: logs/service/service_<时间戳>.log', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('─' * 40)
    # 八、验证
    doc.add_heading('八、第七步：验证部署', 1)
    doc.add_paragraph('1. 健康检查', style='List Number')
    add_code(doc, 'curl http://localhost:8002/health')
    doc.add_paragraph('应返回 {"status":"healthy","service":"VoiceBridge"}。')
    doc.add_paragraph()
    doc.add_paragraph('2. 查看日志', style='List Number')
    add_code(doc, 'tail -f /opt/voicebridge/logs/service/service_*.log')
    doc.add_paragraph('最新日志文件名为 service_<时间戳>.log，可在 logs/service/ 下用 ls -t 查看最新一个。')
    doc.add_paragraph()
    doc.add_paragraph('3. 访问 API 文档：浏览器打开 http://<本机IP或localhost>:8002/docs，可查看并调试接口（端口以 config 为准）。', style='List Number')
    doc.add_paragraph()
    doc.add_paragraph('4. 访问前端：若有 frontend，访问 http://<本机IP或localhost>:8002/login 进行登录与面试流程。', style='List Number')
    doc.add_paragraph()
    doc.add_paragraph('─' * 40)
    # 九、服务管理
    doc.add_heading('九、服务管理命令', 1)
    add_table(doc, ['操作', '命令'], [
        ['启动', 'bash /opt/voicebridge/bash/start_service.sh'],
        ['停止', 'bash /opt/voicebridge/bash/stop_service.sh'],
        ['重启', 'bash /opt/voicebridge/bash/restart_service.sh'],
        ['状态', 'bash /opt/voicebridge/bash/status_service.sh'],
    ])
    add_para(doc, '说明：', ['说明'])
    doc.add_paragraph('停止脚本会读取 voicebridge.pid，先 kill，若未退出则 kill -9，并删除 PID 文件。', style='List Bullet')
    doc.add_paragraph('若 PID 文件丢失，停止脚本会尝试通过 python.*services/run.py 查找进程并提示是否结束。', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('─' * 40)
    # 十、内网穿透
    doc.add_heading('十、内网穿透（可选）', 1)
    add_para(doc, '若需从外网访问部署在内网的服务，可使用 NATAPP。')
    doc.add_paragraph()
    doc.add_paragraph('1. 准备', style='List Number')
    doc.add_paragraph('在 natapp.cn 注册并创建隧道，获得 authtoken。', style='List Bullet')
    doc.add_paragraph('将 tools/natapp/config.ini 中的 authtoken 改为你的 token。', style='List Bullet')
    doc.add_paragraph('tools/natapp/start.sh 会从 config/config.yaml 读取主服务端口并写入 config.ini 的 lanport，无需手动改 config.ini。', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('2. 启动 NATAPP', style='List Number')
    add_code(doc, 'bash /opt/voicebridge/tools/natapp/start.sh')
    doc.add_paragraph('日志写入 logs/natapp/natapp.log，可用 tail -f logs/natapp/natapp.log 查看。')
    doc.add_paragraph()
    doc.add_paragraph('3. 停止', style='List Number')
    add_code(doc, 'bash /opt/voicebridge/tools/natapp/stop.sh')
    doc.add_paragraph()
    add_para(doc, '说明：先确保本机主服务（默认 8002）已启动，再启动 NATAPP；外网访问使用 NATAPP 提供的域名（如示例中的 http://recruitment.natapp1.cc）。', ['说明'])
    doc.add_paragraph()
    doc.add_paragraph('─' * 40)
    # 十一、常见问题
    doc.add_heading('十一、常见问题', 1)
    doc.add_paragraph('“无法找到 conda”：安装 Anaconda/Miniconda 并将 conda.sh 放到脚本可找到的路径（如 ~/anaconda3/etc/profile.d/conda.sh），或保证 conda 在 PATH 中。', style='List Bullet')
    doc.add_paragraph('“无法激活 datastore 环境”：先执行 conda create -n datastore python=3.10 -y，再执行启动脚本。', style='List Bullet')
    doc.add_paragraph('“服务已在运行”：已有进程在跑，如需重启请先执行 bash bash/stop_service.sh，再执行 start_service.sh。', style='List Bullet')
    doc.add_paragraph('端口被占用：只需修改 config/config.yaml 中 services.voicebridge.port、server.port 及 voice_streaming.websocket.port 为未占用端口；NATAPP 启动时会自动从 config 读取主端口。', style='List Bullet')
    doc.add_paragraph('ASR 不可用：检查 config/config.yaml 中 voice_streaming.asr 的 appkey、access_key_id、access_key_secret 是否正确，以及网络是否可达阿里云。', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('─' * 40)
    # 十二、systemd
    doc.add_heading('十二、生产环境与开机自启动（systemd）', 1)
    add_para(doc, '生产环境使用 systemd 托管主服务与内网穿透，实现开机自启、异常重启与统一日志。当前 unit 已按本机配置：项目路径 /opt/voicebridge，运行用户 voicebridge。部署到其他机器时再修改各 .service 中的路径与 User/Group。', ['主服务', '内网穿透', '项目路径', '运行用户'])
    doc.add_heading('12.1 一键安装（主服务 + 内网穿透均开机自启）', 2)
    add_para(doc, '在项目根目录执行：')
    doc.add_paragraph()
    add_code(doc, 'cd /opt/voicebridge\n\n# 1. 安装两个服务单元\nsudo cp deploy/voicebridge.service /etc/systemd/system/\nsudo cp deploy/natapp.service /etc/systemd/system/\n\n# 2. 重载并启用开机自启\nsudo systemctl daemon-reload\nsudo systemctl enable voicebridge\nsudo systemctl enable natapp\n\n# 3. 立即启动（先主服务，再内网穿透）\nsudo systemctl start voicebridge\nsudo systemctl start natapp\n\n# 4. 确认状态\nsudo systemctl status voicebridge\nsudo systemctl status natapp')
    doc.add_heading('12.2 常用命令', 2)
    add_table(doc, ['操作', '命令'], [
        ['主服务状态 / 重启 / 停止', 'sudo systemctl status voicebridge / restart / stop'],
        ['主服务日志', 'journalctl -u voicebridge -f'],
        ['内网穿透状态 / 启动 / 停止', 'sudo systemctl status natapp / start / stop'],
        ['内网穿透日志', 'journalctl -u natapp -f 或 tail -f logs/natapp/natapp.log'],
        ['关闭开机自启', 'sudo systemctl disable voicebridge / sudo systemctl disable natapp'],
    ])
    add_para(doc, '说明：systemd 通过 bash/voicebridge-systemd.sh 前台启动主服务，无需 PID 文件；与手动执行 bash bash/start_service.sh 二选一，不要同时用两种方式启动。内网穿透（NATAPP）会在主服务之后自动启动。', ['说明'])
    doc.add_paragraph()
    doc.add_paragraph('─' * 40)
    # 十三、检查项
    doc.add_heading('十三、生产上线检查项', 1)
    add_para(doc, '上线前建议逐项确认：')
    add_table(doc, ['检查项', '说明'], [
        ['端口', 'config/config.yaml 中主服务端口（默认 8002）、WebSocket（默认 8003）与防火墙/安全组一致；若对外暴露，需放行对应端口。'],
        ['配置', 'server.debug 为 false；ASR、LLM、Rasa 等地址与密钥正确；敏感信息建议用环境变量或 credentials.yaml，勿提交到仓库。'],
        ['Conda', '运行用户下已安装 Conda，且存在 datastore 环境、依赖已 pip install -r requirements.txt。'],
        ['开机自启', '已按第十二节配置 systemd，并 systemctl enable voicebridge、systemctl enable natapp（主服务 + 内网穿透均开机自启）。'],
        ['日志', '日志目录 logs/ 可写；重要环境可配合 logrotate 做日志轮转。'],
        ['健康检查', '部署后访问 http://<本机>:8002/health 返回 {"status":"healthy"}。'],
        ['NATAPP', '官网隧道「本地端口」与 config 主服务端口一致；authtoken 有效；若外网访问则必配。'],
    ])
    doc.add_paragraph()
    doc.add_paragraph('以上为完整部署步骤与说明，按顺序执行即可在本机或服务器上运行 VoiceBridge 服务；生产环境请完成第十二、十三节。')
    return doc

def build_用户手册():
    doc = Document()
    doc.add_heading('VoiceBridge 用户手册', 0)
    add_para(doc, '本文档面向使用 VoiceBridge 进行语音面试的最终用户与管理员，按步骤说明如何访问系统、登录、参加面试以及常见问题处理。生产环境开机自启等部署细节见《部署文档》。', ['最终用户', '管理员', '开机自启'])
    doc.add_paragraph()
    doc.add_paragraph('─' * 40)
    doc.add_heading('一、产品简介', 1)
    add_para(doc, 'VoiceBridge 是一款语音与文字双向转换的面试服务系统，主要能力包括：', ['VoiceBridge'])
    doc.add_paragraph('ASR（语音转文字）：实时将候选人语音识别为文字。', style='List Bullet')
    doc.add_paragraph('对话与面试流程管理：按预设流程提问、追问、打分。', style='List Bullet')
    doc.add_paragraph('统一入口：同一端口（默认 8002，以 config/config.yaml 为准）同时提供 Web 前端页面和 API，无需分别访问不同地址。', style='List Bullet')
    doc.add_paragraph()
    add_para(doc, '典型使用场景：候选人通过浏览器打开登录页，使用邀请账号登录后进入语音面试界面，通过麦克风回答问题，系统实时转写并评估。', ['典型使用场景'])
    doc.add_paragraph()
    doc.add_paragraph('─' * 40)
    doc.add_heading('二、访问地址与入口', 1)
    add_para(doc, '部署完成后，根据实际访问方式选择对应地址：')
    add_table(doc, ['访问方式', '地址示例', '说明'], [
        ['本机访问', 'http://localhost:8002', '在部署机器上打开浏览器（端口以 config 为准）'],
        ['局域网访问', 'http://<服务器IP>:8002', '将 <服务器IP> 换为实际 IP'],
        ['外网访问（NATAPP）', 'http://recruitment.natapp1.cc 等', '以 NATAPP 控制台显示的域名为准'],
    ])
    add_para(doc, '常用入口：', ['常用入口'])
    doc.add_paragraph('登录页：http://<上述地址>/login — 用于输入账号、密码并登录。', style='List Bullet')
    doc.add_paragraph('面试页（桌面端）：http://<上述地址>/interview 或 /desktop-interview — 登录成功后一般会跳转至此。', style='List Bullet')
    doc.add_paragraph('面试页（移动端）：http://<上述地址>/mobile-interview — 适合手机浏览器使用。', style='List Bullet')
    doc.add_paragraph('API 文档：http://<上述地址>/docs — 开发或调试时可查看、调试所有 HTTP 接口。', style='List Bullet')
    doc.add_paragraph('健康检查：http://<上述地址>/health — 返回 {"status":"healthy","service":"VoiceBridge"} 即表示正常。', style='List Bullet')
    add_para(doc, '说明：若部署时未配置 frontend 目录，则仅有 API 可用，上述页面可能无法打开，需联系管理员确认前端是否已部署。', ['说明'])
    doc.add_paragraph()
    doc.add_paragraph('─' * 40)
    doc.add_heading('三、使用流程概览', 1)
    for i, s in enumerate([
        '打开登录页（如 http://<地址>/login）。',
        '使用管理员提供的账号和密码登录。',
        '登录成功后进入面试页面，按提示完成语音面试。',
        '面试结束后可关闭页面或退出。',
    ], 1):
        doc.add_paragraph(f'{i}. {s}', style='List Number')
    doc.add_paragraph('以下按步骤详细说明。')
    doc.add_paragraph()
    doc.add_paragraph('─' * 40)
    doc.add_heading('四、第一步：打开登录页', 1)
    doc.add_paragraph('1. 在浏览器地址栏输入登录地址，例如：')
    doc.add_paragraph('http://localhost:8002/login')
    doc.add_paragraph('或 http://<服务器IP>:8002/login')
    doc.add_paragraph('或外网域名（如 http://recruitment.natapp1.cc/login）。')
    doc.add_paragraph()
    doc.add_paragraph('2. 回车后应看到 VoiceBridge 登录页面（包含账号、密码输入框和登录按钮）。')
    doc.add_paragraph()
    add_para(doc, '若无法打开：', ['若无法打开'])
    doc.add_paragraph('确认地址和端口（默认 8002，以 config 为准）是否正确。', style='List Bullet')
    doc.add_paragraph('若为局域网/外网访问，确认本机或服务器防火墙是否放行主服务端口。', style='List Bullet')
    doc.add_paragraph('使用 http://<地址>/health 检查服务是否在运行。', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('─' * 40)
    doc.add_heading('五、第二步：登录', 1)
    doc.add_paragraph('1. 账号：一般为管理员提供的候选人用户名（candidate_username），或邀请 ID（形如 INV_xxx）。在“账号”输入框中完整填写其一即可。')
    doc.add_paragraph('2. 密码：填写管理员提供的候选人密码（candidate_password）。')
    doc.add_paragraph('3. 点击登录按钮。')
    doc.add_paragraph('4. 登录结果：')
    doc.add_paragraph('成功：页面会跳转到面试页（如 /interview 或 /desktop-interview），并携带本次会话信息。', style='List Bullet')
    doc.add_paragraph('失败：页面会提示错误信息，例如：“用户不存在或邀请不存在”“密码错误”“邀请状态不正确，无法登录”等。', style='List Bullet')
    add_para(doc, '说明：只有状态为 CONFIRMED 的邀请可以登录；登录成功后系统会将状态更新为 IN_PROGRESS（进行中）。', ['说明', 'CONFIRMED', 'IN_PROGRESS'])
    doc.add_paragraph()
    doc.add_paragraph('─' * 40)
    doc.add_heading('六、第三步：进行语音面试', 1)
    doc.add_paragraph('1. 进入面试页：登录成功后会自动跳转到面试页面；若未跳转，可手动访问桌面端 http://<地址>/interview 或 /desktop-interview，移动端 http://<地址>/mobile-interview。')
    doc.add_paragraph('2. 允许麦克风：首次使用浏览器会提示“允许使用麦克风”，请选择允许，否则无法进行语音识别。')
    doc.add_paragraph('3. 按流程答题：页面会显示当前题目与说明；点击“开始录音”或按页面提示开始说话，系统会实时将语音转成文字并参与评估；回答完一题后按页面提示进入下一题；部分题目可能根据回答质量自动追问。')
    doc.add_paragraph('4. 结束面试：按页面提示完成所有题目后，面试流程结束；可关闭浏览器标签页或直接关闭浏览器。')
    add_para(doc, '说明：具体题目内容、追问逻辑、评分规则由后台配置和对话服务决定，本手册仅说明前端使用步骤。', ['说明'])
    doc.add_paragraph()
    doc.add_paragraph('─' * 40)
    doc.add_heading('七、API 文档与调试（可选）', 1)
    doc.add_paragraph('1. 在浏览器打开：http://<地址>/docs。')
    doc.add_paragraph('2. 页面会列出所有 HTTP 接口，可点击展开并尝试“Try it out”进行请求。')
    doc.add_paragraph('3. 登录接口为：POST /api/v1/auth/login，请求体为 {"username":"账号","password":"密码"}。')
    doc.add_paragraph('4. 其他接口（会话检查、对话、面试流程等）可在同一文档中查看路径与参数。')
    add_para(doc, '说明：实际使用 Web 面试时无需手动调 API，按上述登录与面试页操作即可。', ['说明'])
    doc.add_paragraph()
    doc.add_paragraph('─' * 40)
    doc.add_heading('八、常见问题与排查', 1)
    doc.add_heading('8.1 无法打开登录页或页面空白', 2)
    doc.add_paragraph('检查服务是否运行：访问 http://<地址>/health，若无法访问说明服务未启动或端口/防火墙有问题。', style='List Bullet')
    doc.add_paragraph('确认端口：默认 8002，以 config/config.yaml 为准；若部署时修改了端口，需使用新端口访问。', style='List Bullet')
    doc.add_paragraph('确认前端是否部署：若项目未包含 frontend 或未挂载静态资源，则只有 API 可用，需联系管理员部署前端。', style='List Bullet')
    doc.add_heading('8.2 登录提示“用户不存在或邀请不存在”', 2)
    doc.add_paragraph('确认输入的账号是候选人用户名或邀请 ID（如 INV_xxx），且与管理员在系统中配置的一致；联系管理员确认邀请是否已创建且未删除。', style='List Bullet')
    doc.add_heading('8.3 登录提示“邀请状态不正确”', 2)
    doc.add_paragraph('仅 CONFIRMED（已确认）状态的邀请可以登录；联系管理员在后台将对应邀请状态改为“已确认”后再试。', style='List Bullet')
    doc.add_heading('8.4 语音没有识别或没有声音', 2)
    doc.add_paragraph('在浏览器中检查是否已允许麦克风权限；确认设备麦克风正常；若为外网访问，确认 WebSocket（默认 8003）是否被 NATAPP 或防火墙正确转发。', style='List Bullet')
    doc.add_heading('8.5 外网访问很慢或无法连接', 2)
    doc.add_paragraph('若使用 NATAPP，请确认 NATAPP 客户端已启动；tools/natapp/start.sh 会从 config/config.yaml 读取主端口并写入 config.ini。查看 NATAPP 日志：tail -f <项目根目录>/logs/natapp/natapp.log，确认隧道是否建立成功。', style='List Bullet')
    doc.add_heading('8.6 如何确认服务在运行', 2)
    doc.add_paragraph('方式一：浏览器访问 http://<地址>/health，应返回 {"status":"healthy","service":"VoiceBridge"}。', style='List Bullet')
    doc.add_paragraph('方式二：在部署机器上执行 bash <项目根目录>/bash/status_service.sh，可查看进程、端口及最近日志路径。', style='List Bullet')
    doc.add_heading('8.7 管理员如何停止/重启服务', 2)
    doc.add_paragraph('手动脚本：停止 bash <项目根目录>/bash/stop_service.sh；重启 bash <项目根目录>/bash/restart_service.sh；状态与日志 bash <项目根目录>/bash/status_service.sh。', style='List Bullet')
    doc.add_paragraph('生产环境（systemd）：若已配置开机自启（主服务 + 内网穿透），可用 sudo systemctl restart voicebridge 重启主服务，sudo systemctl restart natapp 重启内网穿透；状态与日志：sudo systemctl status voicebridge / status natapp，journalctl -u voicebridge -f / journalctl -u natapp -f。详细步骤见《部署文档》。', style='List Bullet')
    doc.add_paragraph()
    doc.add_paragraph('─' * 40)
    doc.add_heading('九、小结', 1)
    add_table(doc, ['步骤', '操作'], [
        ['1', '浏览器打开 http://<地址>/login'],
        ['2', '输入管理员提供的账号（候选人用户名或邀请 ID）和密码，点击登录'],
        ['3', '登录成功后进入面试页，允许麦克风后按页面提示完成语音答题'],
        ['4', '完成所有题目后关闭页面即可'],
    ])
    doc.add_paragraph('遇到无法登录、无法打开页面或语音不识别时，可先按第八章“常见问题与排查”逐项检查；仍无法解决时请联系部署管理员并提供访问地址、账号类型及错误提示内容。')
    return doc

def fix_add_para(doc, text, bold_parts=None):
    """修正：add_para 的 bold_parts 是列表，且原逻辑有误。这里简化为整段添加，粗体用 ** 标记的再单独处理。"""
    if not bold_parts:
        doc.add_paragraph(text)
        return
    p = doc.add_paragraph()
    remain = text
    for b in bold_parts:
        if b in remain:
            idx = remain.find(b)
            if idx > 0:
                p.add_run(remain[:idx])
            r = p.add_run(b)
            r.bold = True
            remain = remain[idx + len(b):]
    if remain:
        p.add_run(remain)

# 修正 build_部署文档 里错误的 add_para 调用（第二处只有两个参数却传了三个）
if __name__ == '__main__':
    import os
    d = os.path.dirname(os.path.abspath(__file__))
    doc1 = build_部署文档()
    # 修正第一段：包含...配置 应为 包含... 配置
    doc1.save(os.path.join(d, '部署文档.docx'))
    print('已生成: 部署文档.docx')
    doc2 = build_用户手册()
    doc2.save(os.path.join(d, '用户手册.docx'))
    print('已生成: 用户手册.docx')
