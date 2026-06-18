# -*- coding: utf-8 -*-
"""
从Word文档中提取各段落关键点，生成新的Word文档
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def extract_key_points(text):
    """
    从文本中提取关键点
    使用智能规则：提取重要信息、数字、结论等
    """
    key_points = []
    
    # 移除多余空白
    text = re.sub(r'\s+', ' ', text.strip())
    
    # 如果文本太短，直接返回
    if len(text) < 10:
        return []
    
    # 如果文本较短（小于150字），直接返回
    if len(text) < 150:
        return [text]
    
    # 提取包含数字、百分比、指标的关键信息
    important_keywords = [
        # 技术指标
        '准确率', '延迟', '并发', '性能', '效率', '响应时间', '吞吐量',
        # 功能特性
        '实现', '完成', '支持', '达到', '提升', '优化', '改进', '新增',
        # 系统相关
        '功能', '系统', '架构', '技术', '方案', '问题', '解决', '部署',
        # 业务相关
        '面试', '评价', '评分', '追问', '转写', '识别', '语音',
        # 总结性词汇
        '总结', '结论', '要点', '关键', '核心', '主要', '重要'
    ]
    
    # 按句号、感叹号、问号、分号分割
    sentences = re.split(r'[。！？；\n]', text)
    
    for sentence in sentences:
        sentence = sentence.strip()
        if len(sentence) < 8:  # 太短的句子跳过
            continue
        
        # 检查是否包含重要关键词
        has_keyword = any(keyword in sentence for keyword in important_keywords)
        
        # 检查是否包含数字（百分比、数量等）
        has_number = bool(re.search(r'\d+[%％]|\d+\.\d+%|达到\d+|超过\d+|低于\d+', sentence))
        
        # 检查是否是总结性语句（通常包含"总之"、"综上所述"等）
        is_summary = bool(re.search(r'总之|综上所述|总的来说|总而言之|简而言之', sentence))
        
        # 如果满足任一条件，添加到关键点
        if has_keyword or has_number or is_summary:
            # 限制句子长度，避免过长
            if len(sentence) > 200:
                sentence = sentence[:200] + '...'
            key_points.append(sentence)
    
    # 如果找到了关键点，返回
    if key_points:
        # 去重并限制数量
        seen = set()
        unique_points = []
        for point in key_points:
            if point not in seen:
                seen.add(point)
                unique_points.append(point)
                if len(unique_points) >= 5:  # 最多5个关键点
                    break
        return unique_points
    
    # 如果没有找到关键点，提取前两句话作为摘要
    sentences = re.split(r'[。！？]', text)
    summary = []
    for sentence in sentences[:2]:
        sentence = sentence.strip()
        if len(sentence) >= 10:
            summary.append(sentence)
            if len(summary) >= 2:
                break
    
    return summary if summary else [text[:150] + '...' if len(text) > 150 else text]


def process_document(input_path, output_path):
    """
    处理Word文档，提取关键点并生成新文档
    """
    print(f"正在读取文档: {input_path}")
    doc = Document(input_path)
    
    # 创建新文档
    new_doc = Document()
    
    # 设置标题
    title = new_doc.add_heading('招聘智能体项目汇报文档 - 关键点提炼', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加说明
    note = new_doc.add_paragraph('本文档从原文档中提取了各段落的关键信息点')
    note_format = note.runs[0].font
    note_format.italic = True
    note_format.size = Pt(10)
    try:
        note_format.color.rgb = RGBColor(128, 128, 128)
    except:
        pass  # 如果颜色设置失败，忽略
    
    new_doc.add_paragraph()  # 空行
    
    section_count = 0
    paragraph_count = 0
    last_was_heading = False
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # 跳过空段落
        if not text:
            continue
        
        # 检查是否是标题（通常标题较短且格式特殊）
        is_heading = False
        
        # 检查样式名称
        style_name = para.style.name.lower()
        if 'heading' in style_name or '标题' in style_name or 'heading' in para.style.name:
            is_heading = True
        
        # 检查格式：粗体且较短
        if not is_heading and para.runs:
            is_bold = any(run.bold for run in para.runs if run.bold)
            if is_bold and len(text) < 80:
                is_heading = True
        
        # 检查内容模式：数字开头（如"1. "、"一、"等）
        if not is_heading:
            if re.match(r'^[\d一二三四五六七八九十]+[\.、．]', text) or \
               re.match(r'^第[一二三四五六七八九十\d]+[章节部分]', text):
                is_heading = True
        
        # 如果是标题，添加为标题格式
        if is_heading:
            section_count += 1
            heading = new_doc.add_heading(text, level=1)
            paragraph_count = 0
            last_was_heading = True
        else:
            # 提取关键点
            key_points = extract_key_points(text)
            
            if key_points:
                # 如果上一段是标题，添加空行
                if last_was_heading:
                    new_doc.add_paragraph()
                    last_was_heading = False
                
                paragraph_count += 1
                # 添加段落编号（只在有标题的情况下显示）
                if section_count > 0:
                    para_text = f"【{section_count}.{paragraph_count}】"
                else:
                    para_text = f"【{paragraph_count}】"
                
                para_heading = new_doc.add_paragraph(para_text)
                para_heading.runs[0].bold = True
                para_heading.runs[0].font.size = Pt(11)
                
                # 添加关键点
                for point in key_points:
                    point_text = f"  • {point}"
                    point_para = new_doc.add_paragraph(point_text)
                    point_para.runs[0].font.size = Pt(10.5)
                    point_para.paragraph_format.first_line_indent = Pt(21)  # 首行缩进
                
                new_doc.add_paragraph()  # 段落间空行
    
    # 处理表格
    for table in doc.tables:
        new_doc.add_paragraph()  # 空行
        table_heading = new_doc.add_heading('表格内容', level=2)
        
        # 创建新表格
        new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        new_table.style = 'Light Grid Accent 1'
        
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_table.rows[i].cells[j].text = cell.text.strip()
    
    # 保存文档
    print(f"正在保存文档: {output_path}")
    new_doc.save(output_path)
    print(f"文档已成功生成: {output_path}")


if __name__ == '__main__':
    # 输入和输出路径
    input_file = Path(__file__).parent.parent / 'docs' / '招聘智能体项目汇报文档(汇总版)_20260209.docx'
    output_file = Path(__file__).parent.parent / 'docs' / '招聘智能体项目汇报文档-关键点提炼.docx'
    
    if not input_file.exists():
        print(f"错误: 找不到输入文件 {input_file}")
        exit(1)
    
    try:
        process_document(input_file, output_file)
        print("处理完成！")
    except Exception as e:
        print(f"处理过程中出现错误: {e}")
        import traceback
        traceback.print_exc()

